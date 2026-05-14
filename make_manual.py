from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── Color constants ─────────────────────────────────────────────
RED      = RGBColor(0xDC, 0x14, 0x3C)
DARK     = RGBColor(0x08, 0x0E, 0x1C)
GOLD     = RGBColor(0xD4, 0xA0, 0x17)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GRAY     = RGBColor(0x64, 0x74, 0x8B)
LIGHTBG  = RGBColor(0xF8, 0xFA, 0xFC)
STEPBG   = RGBColor(0xFF, 0xF0, 0xF3)
GREENBG  = RGBColor(0xF0, 0xFF, 0xF4)
BLUEBG   = RGBColor(0xEF, 0xF6, 0xFF)

# ── Helper: set paragraph shading ───────────────────────────────
def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, border_color='DC143C', size=8):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), border_color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

# ── Helper: add a run to a paragraph ────────────────────────────
def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run

# ── Cover Page ──────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(cover, '080E1C')
run = cover.add_run('\n\n🇰🇳  CAMPAIGN 365\nSKNLP — St. Kitts-Nevis Labour Party\n\n')
run.bold = True
run.font.color.rgb = WHITE
run.font.size = Pt(28)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(sub, 'DC143C')
r = sub.add_run('   USER GUIDE — Web Dashboard & Mobile App   ')
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(16)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
n = note.add_run('\nThis guide will walk you through every step, one at a time.\nNo technical knowledge required.\n')
n.italic = True
n.font.color.rgb = GRAY
n.font.size = Pt(11)

doc.add_paragraph()

# ── Login Credentials Box ────────────────────────────────────────
cred_title = doc.add_paragraph()
set_paragraph_shading(cred_title, '1E293B')
r = cred_title.add_run('  🔑  YOUR LOGIN CREDENTIALS  ')
r.bold = True
r.font.color.rgb = GOLD
r.font.size = Pt(13)

creds = [
    ('Web Dashboard Email', 'admin@sknlp.campaign365.app'),
    ('Password',            'password'),
    ('Mobile App Email',    'admin@sknlp.campaign365.app'),
    ('Mobile App Password', 'password'),
]
for label, val in creds:
    p = doc.add_paragraph()
    set_paragraph_shading(p, 'F1F5F9')
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, f'  {label}:  ', bold=True, color=DARK, size=11)
    add_run(p, val, bold=False, color=RED, size=11)

tip = doc.add_paragraph()
set_paragraph_shading(tip, 'FFF8E1')
add_run(tip, '  💡 Tip: ', bold=True, color=GOLD, size=10)
add_run(tip, 'Save these credentials somewhere safe. You can change your password after first login.', size=10)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 1 — WEB DASHBOARD
# ════════════════════════════════════════════════════════════════
part1 = doc.add_paragraph()
part1.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(part1, 'DC143C')
r = part1.add_run('\n  PART 1 — WEB DASHBOARD  \n')
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(20)

doc.add_paragraph()

def section_header(text, color_hex='DC143C', text_color=WHITE):
    p = doc.add_paragraph()
    set_paragraph_shading(p, color_hex)
    r = p.add_run(f'  {text}  ')
    r.bold = True
    r.font.color.rgb = text_color
    r.font.size = Pt(13)
    return p

def step_box(number, title, steps, tip_text=None, warning_text=None):
    # Step header
    hdr = doc.add_paragraph()
    set_paragraph_shading(hdr, '1E293B')
    hdr.paragraph_format.space_before = Pt(6)
    r = hdr.add_run(f'  STEP {number}  ')
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)
    r2 = hdr.add_run(f'  {title}')
    r2.bold = True
    r2.font.color.rgb = WHITE
    r2.font.size = Pt(12)

    # Steps
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        set_paragraph_shading(p, 'F8FAFC')
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        add_run(p, f'  {i}.  ', bold=True, color=RED, size=11)
        # handle bold parts wrapped in **
        parts = step.split('**')
        for j, part in enumerate(parts):
            if j % 2 == 1:
                add_run(p, part, bold=True, color=DARK, size=11)
            else:
                add_run(p, part, bold=False, color=DARK, size=11)

    if tip_text:
        tp = doc.add_paragraph()
        set_paragraph_shading(tp, 'ECFDF5')
        tp.paragraph_format.left_indent = Cm(0.5)
        add_run(tp, '  ✅  ', bold=True, size=11)
        add_run(tp, tip_text, italic=True, color=RGBColor(0x05, 0x96, 0x4B), size=10)

    if warning_text:
        wp = doc.add_paragraph()
        set_paragraph_shading(wp, 'FFF3CD')
        wp.paragraph_format.left_indent = Cm(0.5)
        add_run(wp, '  ⚠️  ', bold=True, size=11)
        add_run(wp, warning_text, italic=True, color=RGBColor(0x92, 0x40, 0x00), size=10)

    doc.add_paragraph()

# ── STEP 1: Open Dashboard ───────────────────────────────────────
step_box(1, 'Open the Web Dashboard', [
    'Open your internet browser (Chrome or Firefox works best).',
    'Type the website address that was given to you into the address bar at the top.',
    'Press **Enter** on your keyboard.',
    'You will see the SKNLP login page appear on your screen.',
])

# ── STEP 2: Log In ───────────────────────────────────────────────
step_box(2, 'Log In to Your Account', [
    'On the login page, click on the **Email** box.',
    'Type: **admin@sknlp.campaign365.app**',
    'Click on the **Password** box.',
    'Type: **password**',
    'Click the **Log In** button.',
    'Wait a moment — the system will check your details.',
    'You will be taken to the **Dashboard** (the main home screen).',
], tip_text='If you see an error, double-check that you typed the email and password exactly as shown above.')

# ── STEP 3: Dashboard ────────────────────────────────────────────
step_box(3, 'Understanding the Dashboard (Home Screen)', [
    'After logging in, you will see the **Dashboard** — this is your home screen.',
    'At the top left, you will see the **SKNLP logo** and your party name.',
    'On the **left side** of the screen, there is a menu (called the Sidebar) with links to all sections.',
    'In the **middle of the screen**, you will see 4 number boxes called Stat Cards:',
    '   — **Voters Contacted:** How many voters your team has spoken to.',
    '   — **Total Voters:** Total number of voters in the system.',
    '   — **Active Campaigns:** How many campaigns are running right now.',
    '   — **Canvassers Online:** How many field workers are active right now.',
    'These numbers update automatically. Press **F5** on your keyboard to refresh and see the latest.',
])

# ── STEP 4: Voters ───────────────────────────────────────────────
step_box(4, 'View and Manage Voters', [
    'On the left sidebar, click on **"Voters"**.',
    'You will see a table with all voters in your database (180+ voters).',
    'Each row shows the voter\'s name, constituency, support level, and contact info.',
    'To **search** for a specific voter: click the search box at the top and type their name.',
    'To **filter voters** by constituency or support level: click the filter dropdowns above the table.',
    'To **view one voter\'s full profile**: click their name in the table.',
    '   — Their profile page shows: address, phone, age, voting history, and support level.',
    'To **add a new voter manually**: click the red **"Add Voter"** button at the top right.',
    '   — Fill in their details and click **Save**.',
], tip_text='Support levels are: Supporter (they support SKNLP), Undecided (not sure), Opposition (they do not support).')

# ── STEP 5: Campaigns ────────────────────────────────────────────
step_box(5, 'View Campaigns', [
    'On the left sidebar, click on **"Campaigns"**.',
    'You will see a list of all campaigns — Active, Upcoming, and Completed.',
    'The current main campaign is: **2026 General Election — National (Active)**.',
    'Click on any campaign name to open it and see more details.',
    '   — Details include: number of voters targeted, how many have responded, and campaign progress.',
])

# ── STEP 6: Canvassing ───────────────────────────────────────────
step_box(6, 'View Field Canvassing Progress', [
    'On the left sidebar, click on **"Canvassing"**.',
    'You will see a map and a list of Walk Lists (canvassing zones).',
    'There are 3 current zones:',
    '   — **St. Christopher North — Zone A:** 88% complete (280 out of 320 doors knocked)',
    '   — **St. Christopher North — Zone B:** 100% complete (Finished ✅)',
    '   — **Basseterre Central Walk:** 44% complete (180 out of 410 doors knocked)',
    'The progress bars show how much of each zone is done.',
    'You can see which canvasser is assigned to each zone.',
])

# ── STEP 7: Team ─────────────────────────────────────────────────
step_box(7, 'Manage Your Team', [
    'On the left sidebar, click on **"Team"**.',
    'You will see a list of all team members — Canvassers, Runners, Outdoor Agents.',
    'To **add a new team member**: click the red **"Add Member"** button.',
    '   — Enter their name, email, phone number, and select their role.',
    '   — Click **Save** — they will receive login details.',
    'To **view a team member\'s details**: click their name.',
    'To **remove a team member**: click on their name and select **"Remove"**.',
])

# ── STEP 8: Communications ────────────────────────────────────────
step_box(8, 'Send a Message to Voters (Email or SMS)', [
    'On the left sidebar, click on **"Communications"**.',
    'To send a new message, click the red **"New Campaign"** button.',
    'Choose your message type: **Email** or **SMS**.',
    'In the **Subject** box, type the subject of your message.',
    'In the **Message** box, type what you want to say.',
    'Under **Audience**, select who should receive it:',
    '   — All Voters / Supporters Only / Undecided Voters.',
    'Click **"Send Now"** to send immediately.',
    'Your sent messages will appear in the list below.',
], tip_text='SMS messages go directly to voters\' phones. Email goes to their email address.')

# ── STEP 9: Reports ───────────────────────────────────────────────
step_box(9, 'View Reports and Statistics', [
    'On the left sidebar, click on **"Reports"**.',
    'You will see charts and numbers showing your campaign progress.',
    'Key numbers shown:',
    '   — **Voter Contact Rate:** What percentage of voters have been contacted.',
    '   — **Turnout Projection:** What percentage are expected to vote.',
    '   — **Canvassing Efficiency:** How productive your canvassing team is.',
    '   — **Fundraising Raised:** Total money raised so far.',
    'To **download a report**: click the **"Export"** button and choose PDF or Excel.',
])

# ── STEP 10: Fundraising ─────────────────────────────────────────
step_box(10, 'View and Manage Fundraising', [
    'On the left sidebar, click on **"Fundraising"**.',
    'You will see the total amount raised vs. your fundraising goal.',
    'The bar at the top shows your progress towards the goal.',
    'Below that, you can see your **Top Donors** list.',
    'To **add a new donation**: click **"Add Donation"** button.',
    '   — Enter the donor\'s name and donation amount.',
    '   — Click **Save**.',
    'A chart at the bottom shows donation trends over time.',
])

# ── STEP 11: Notifications ────────────────────────────────────────
step_box(11, 'Check Notifications', [
    'Look at the **top right corner** of the screen.',
    'You will see a **bell icon** (🔔) with a number on it.',
    'The number shows how many new notifications you have.',
    'Click the **bell icon** to open your notifications.',
    'Notifications show things like: new donations, canvassing updates, urgent alerts.',
    'Click on any notification to see more details.',
    'Click **"Mark All Read"** to clear all notifications.',
])

# ── STEP 12: Log Out ─────────────────────────────────────────────
step_box(12, 'Log Out of the Dashboard', [
    'Look at the **top right corner** of the screen.',
    'Click on your **name or profile picture**.',
    'A small menu will appear.',
    'Click **"Log Out"**.',
    'You will be taken back to the login page.',
], tip_text='Next time you visit the website, you will stay logged in automatically. You do not need to log in again unless you clicked Log Out.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# PART 2 — MOBILE APP
# ════════════════════════════════════════════════════════════════
part2 = doc.add_paragraph()
part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(part2, '080E1C')
r = part2.add_run('\n  PART 2 — MOBILE APP (Android)  \n')
r.bold = True
r.font.color.rgb = WHITE
r.font.size = Pt(20)

doc.add_paragraph()

# ── STEP 1: Install ──────────────────────────────────────────────
step_box(1, 'Install the App on Your Android Phone', [
    'You will receive a **download link** via WhatsApp or Email.',
    'Open that link on your **Android phone** (not on a computer).',
    'Tap the **"Download APK"** button.',
    'Your phone may show a message: **"Allow install from unknown sources"** — tap **Allow** or **OK**.',
    'Once downloaded, tap the file to **Install** it.',
    'After installing, look for the **Campaign 365** icon on your home screen.',
    'Tap the icon to open the app.',
], warning_text='If your phone does not allow the install, go to Phone Settings → Security → turn on "Install from Unknown Sources", then try again.')

# ── STEP 2: Log In ───────────────────────────────────────────────
step_box(2, 'Log In to the Mobile App', [
    'Open the **Campaign 365** app.',
    'You will see the SKNLP logo and a login screen.',
    'Tap on the **Email** box.',
    'Type: **admin@sknlp.campaign365.app**',
    'Tap on the **Password** box.',
    'Type: **password**',
    'Tap the **"Sign In"** button.',
    '   👉 **Shortcut:** At the bottom of the login screen, tap the **"👤 Admin"** quick-fill button — it fills in your email and password automatically!',
    'Wait a moment — you will be taken to the **Home Screen**.',
], tip_text='Once you log in, the app remembers you. You will not need to log in again unless you log out.')

# ── STEP 3: Home Screen ──────────────────────────────────────────
step_box(3, 'Understanding the Home Screen', [
    'The Home Screen is the main page of the app.',
    'At the **top**, you will see the SKNLP logo and your name.',
    'Below that, there are **3 live stat boxes** showing:',
    '   — **Contacted:** How many voters have been spoken to.',
    '   — **Doors Today:** How many doors were knocked today.',
    '   — **Active Now:** How many canvassers are working right now.',
    'In the **middle**, there is a map of your canvassing area.',
    'Below the map, there are **Turf Cards** — each card shows a walk list with its progress.',
    'At the **very bottom**, there is a row of 5 tabs to navigate the app.',
    'Below the turf cards, there are **Quick Access buttons** — tap any of them to go to that feature.',
])

# ── STEP 4: Canvassing ───────────────────────────────────────────
step_box(4, 'Start Canvassing (Door to Door)', [
    'At the bottom of the screen, tap the **"Canvass"** tab (map icon 🗺️).',
    'You will see a map with your current area highlighted.',
    'At the **bottom of the screen**, you will see the next voter\'s name and address.',
    'Go to that address and knock on the door.',
    'After speaking with the voter, select their **support level**:',
    '   — Tap **"Support"** — if they support SKNLP',
    '   — Tap **"Undecided"** — if they are not sure',
    '   — Tap **"Oppose"** — if they do not support SKNLP',
    'Tap **"Save & Next Voter →"** to record this visit and move to the next voter.',
], tip_text='Every vote you record is automatically saved, even without internet. It will sync when you reconnect.')

# ── STEP 5: Mark Not Home ─────────────────────────────────────────
step_box(5, 'Mark a Voter as "Not Home"', [
    'If you knock and nobody answers, tap the **"Not Home"** button.',
    'The app will record this and move you to the next voter.',
    'You can revisit "Not Home" voters later from your Canvassing History.',
])

# ── STEP 6: Quick Add Voter ──────────────────────────────────────
step_box(6, 'Add a New Voter You Meet at the Door', [
    'If you meet someone who is **not in the system**, you can add them right away.',
    'From the Home Screen, tap the **"➕ Quick Add Voter"** button.',
    'Fill in their details:',
    '   — **Name:** Type their full name.',
    '   — **Address:** Type their home address.',
    '   — **Phone Number:** Type their phone number.',
    '   — **Email:** (Optional) Type their email.',
    'Select their **Support Level** by tapping one of the coloured buttons.',
    'Tap **"Save & Add to Turf"**.',
    'A green tick (✅) will appear confirming the voter was added.',
], tip_text='This voter will now appear in your voter database on the web dashboard as well.')

# ── STEP 7: View Voter Profile ────────────────────────────────────
step_box(7, 'Look Up a Voter\'s Profile', [
    'From the Home Screen, tap **"👤 Voter Profile"**.',
    'You will see a voter\'s full profile page.',
    'It shows their name, address, constituency, support level, and contact history.',
    'Use the **tabs** at the top to switch between Profile, Voting History, and Notes.',
])

# ── STEP 8: Canvassing History ────────────────────────────────────
step_box(8, 'Check Your Canvassing History', [
    'From the Home Screen, tap **"📜 Canvass History"**.',
    'You will see a list of all your past canvassing sessions.',
    'Each card shows:',
    '   — The **date** of the session.',
    '   — The **location** (area name).',
    '   — **Number of voters** contacted in that session.',
    '   — A green tick ✅ if the data has been synced to the cloud.',
    'Tap **"This Week"** or **"All Time"** at the top to filter by time period.',
])

# ── STEP 9: Quick Reports ────────────────────────────────────────
step_box(9, 'View Your Performance Reports', [
    'From the Home Screen, tap **"📊 Quick Reports"**.',
    'You will see 4 boxes with your key statistics:',
    '   — **This Week:** How many doors you knocked this week.',
    '   — **KPI Summary:** How many voters you persuaded.',
    '   — **Support Trend:** A bar chart showing daily support numbers.',
    '   — **Voter Sentiment:** A pie chart showing Supporter / Undecided / Opposition percentages.',
    'Tap **"Export Report"** to share your report with your cluster manager.',
])

# ── STEP 10: GOTV Mode ────────────────────────────────────────────
step_box(10, 'Use GOTV Mode on Election Day', [
    'From the Home Screen, tap **"🗳️ GOTV Mode"**.',
    'Tap the **"Activate GOTV"** toggle to turn it ON.',
    'You will see a list of priority voters who have NOT voted yet.',
    'For each voter, you can:',
    '   — Tap **"Call Now"** to call them directly from your phone.',
    '   — Tap **"Send Reminder"** to send them a text message.',
    'The **heat map** at the bottom shows which areas have high or low voter turnout.',
    'The **Target Box** shows how many more votes you need to hit your goal.',
], tip_text='Use GOTV Mode on Election Day only. It is designed to help you reach voters who have not voted yet.')

# ── STEP 11: Runner Tracking ─────────────────────────────────────
step_box(11, 'Track Your Runners (For Supervisors)', [
    'From the Home Screen, tap **"🏃 Runner Tracking"**.',
    'You will see a live map with coloured dots:',
    '   — 🔵 **Blue dots** = Canvassers',
    '   — 🟡 **Gold dots** = Runners',
    '   — 🔴 **Red dot** = Panic alert location',
    'At the bottom of the screen, you will see a list of runners with their names and current tasks.',
    'Tap **"Contact Runner"** to call a specific runner.',
    'Tap **"Request Backup"** to send help to that location.',
])

# ── STEP 12: AI Insights ──────────────────────────────────────────
step_box(12, 'View AI Insights and Recommendations', [
    'From the Home Screen, tap **"🤖 AI Insights"**.',
    'The AI will show you smart recommendations based on your data:',
    '   — Which voter age groups are most persuadable.',
    '   — Recommended next actions for your team.',
    '   — Overall voter sentiment summary.',
    '   — Turnout predictions for each constituency.',
    'These insights help you focus your efforts where they matter most.',
])

# ── STEP 13: Report an Issue ─────────────────────────────────────
step_box(13, 'Report a Problem or Unsafe Situation', [
    'If you encounter a hostile voter or feel unsafe, tap **"⚠️ Report Issue"** from the Home Screen.',
    'You will see different types of issues to choose from.',
    'Tap the one that matches your situation (e.g. **"Hostile Voter"**).',
    'Add any extra notes in the text box.',
    'Tap **"Submit Report"**.',
    'Your cluster manager will be notified immediately.',
], warning_text='If you are in immediate danger, use the Panic Button (Step 14) instead.')

# ── STEP 14: Panic Button ────────────────────────────────────────
step_box(14, 'Use the Panic Button in an Emergency', [
    'From the Home Screen, tap **"🚨 Panic Screen"**.',
    'A red screen will appear with a **10-second countdown**.',
    'If you do NOT tap **"Cancel"** within 10 seconds, an emergency alert is automatically sent to HQ with your exact GPS location.',
    'To **cancel** (if it was a mistake): tap **"Cancel"** before the countdown reaches zero.',
], warning_text='Only use this in a REAL emergency. A false panic alert will cause your team to respond unnecessarily.')

# ── STEP 15: My Schedule ─────────────────────────────────────────
step_box(15, 'Check Your Schedule', [
    'From the Home Screen, tap **"📅 My Schedule"**.',
    'You will see a weekly calendar at the top with today highlighted.',
    'Tap any day to see the tasks assigned to you for that day.',
    'Each task shows the task name, time, and location.',
    'Tap **"Start Now"** on any task to begin it.',
])

# ── STEP 16: Sync Your Data ─────────────────────────────────────
step_box(16, 'Sync Your Data to the Cloud', [
    'The app automatically saves everything on your phone — even without internet.',
    'When you connect to the internet (Wi-Fi or mobile data), it syncs automatically.',
    'To manually sync: tap the **"Profile"** tab (👤) at the bottom of the screen.',
    'Tap **"Log Out & Sync Data"**.',
    'A loading screen will appear while your data uploads.',
    'When done, you will see a **green checkmark screen** saying "Sync Complete ✨".',
    'Tap **"Back to Home"** to return to the app.',
], tip_text='You will see the number of records uploaded (e.g. 2,341 records). This confirms all your work is safely stored.')

# ── STEP 17: Log Out ──────────────────────────────────────────────
step_box(17, 'Log Out of the Mobile App', [
    'Tap the **"Profile"** tab (👤) at the bottom right of the screen.',
    'Scroll down to the bottom of the profile page.',
    'Tap **"Log Out"**.',
    'A screen will appear asking you to confirm.',
    'Tap **"Log Out & Sync Data"** — this syncs all your data first, then logs you out.',
    'You will be taken back to the login screen.',
], tip_text='Always use "Log Out & Sync Data" so that none of your work is lost.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# FAQ
# ════════════════════════════════════════════════════════════════
section_header('❓  FREQUENTLY ASKED QUESTIONS', '1E293B', WHITE)
doc.add_paragraph()

faqs = [
    (
        'The app says "Cannot connect to server". What do I do?',
        'Check that your phone has internet (mobile data or Wi-Fi). If you have no internet, the app still works — all your data is saved on your phone and will sync later automatically.'
    ),
    (
        'I forgot my password. How do I reset it?',
        'On the login screen, tap "Forgot Password?" and enter your email address. A password reset link will be sent to your email.'
    ),
    (
        'The app is slow or freezing. What should I do?',
        'Close the app completely and reopen it. If the problem continues, restart your phone. Make sure you have enough storage space on your phone.'
    ),
    (
        'Can two people use the same login at the same time?',
        'No. Each person should have their own account. Contact your cluster manager to get separate accounts created for each team member.'
    ),
    (
        'How do I know if my canvassing data has been saved?',
        'After you tap "Save & Next Voter", the app shows the next voter — that means it was saved. You can also check Canvassing History to see all your recorded visits.'
    ),
    (
        'How do I know if my data has synced to the cloud?',
        'Go to Profile → scroll down → you will see "Last Sync" time. A green badge "Offline Mode Ready" means the app is ready to work without internet.'
    ),
    (
        'I added a voter on the mobile app. Will it show on the web dashboard?',
        'Yes, as soon as your phone syncs (when you have internet), the new voter will appear on the web dashboard automatically.'
    ),
    (
        'I accidentally pressed the Panic Button. What do I do?',
        'Tap "Cancel" before the 10-second countdown reaches zero. If you could not cancel in time, immediately call your cluster manager and let them know it was an accident.'
    ),
]

for q, a in faqs:
    qp = doc.add_paragraph()
    set_paragraph_shading(qp, 'FFF0F3')
    qp.paragraph_format.left_indent  = Cm(0.3)
    qp.paragraph_format.space_before = Pt(4)
    add_run(qp, '  Q:  ', bold=True, color=RED, size=11)
    add_run(qp, q, bold=True, color=DARK, size=11)

    ap = doc.add_paragraph()
    set_paragraph_shading(ap, 'F8FAFC')
    ap.paragraph_format.left_indent = Cm(0.3)
    add_run(ap, '  A:  ', bold=True, color=GRAY, size=11)
    add_run(ap, a, color=DARK, size=11)
    doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# Quick Reference Card
# ════════════════════════════════════════════════════════════════
section_header('📱  MOBILE APP — QUICK REFERENCE CARD', 'DC143C', WHITE)
doc.add_paragraph()

ref_title = doc.add_paragraph()
add_run(ref_title, 'All buttons on the Home Screen and what they do:', bold=True, color=DARK, size=11)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
set_cell_shading(hdr_cells[0], '1E293B')
set_cell_shading(hdr_cells[1], '1E293B')
p0 = hdr_cells[0].paragraphs[0]
p1 = hdr_cells[1].paragraphs[0]
add_run(p0, '  Button', bold=True, color=WHITE, size=11)
add_run(p1, '  What it does', bold=True, color=WHITE, size=11)

rows_data = [
    ('📋  Campaign Materials', 'View party posters, flyers, and talking points for voters.'),
    ('🏃  Runner Tracking',    'See all runners\' live GPS locations on a map.'),
    ('🗳️  GOTV Mode',          'Election Day command center — call and remind voters to vote.'),
    ('📊  Quick Reports',      'See your weekly stats, charts, and performance numbers.'),
    ('🤖  AI Insights',        'Smart recommendations powered by AI based on your data.'),
    ('📜  Canvass History',    'See all your past canvassing sessions and sync status.'),
    ('👤  Voter Profile',      'Look up any voter\'s full details and history.'),
    ('🚨  Panic Screen',       'EMERGENCY ONLY — sends SOS alert with your GPS location to HQ.'),
    ('📅  My Schedule',        'View your assigned tasks and schedule for the week.'),
    ('➕  Quick Add Voter',     'Add a new voter you meet at the door, right on the spot.'),
    ('⚠️  Report Issue',        'Report a hostile voter, unsafe situation, or problem.'),
    ('🗳  Election Night',      'Live election night results, constituency by constituency.'),
    ('🤝  Recruit Volunteer',  'Sign up a new volunteer and add them to your team.'),
    ('🚪  Logout Screen',      'Log out and sync all your data to the cloud.'),
    ('✅  Sync Complete',       'Confirmation page showing how many records were uploaded.'),
    ('❓  Help & Training',     'Step-by-step tutorials and frequently asked questions.'),
]

for i, (btn, desc) in enumerate(rows_data):
    row    = table.add_row()
    c0, c1 = row.cells
    fill   = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    set_cell_shading(c0, fill)
    set_cell_shading(c1, fill)
    pb = c0.paragraphs[0]
    pd = c1.paragraphs[0]
    add_run(pb, f'  {btn}', bold=True, color=RED, size=10)
    add_run(pd, f'  {desc}', color=DARK, size=10)

doc.add_paragraph()

# Footer
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(footer_p, '080E1C')
add_run(footer_p, '\n  Campaign 365 — SKNLP Edition   |   Powered by SkillLeo   |   For support, contact your system administrator  \n',
        italic=True, color=GOLD, size=9)

# ── Save ─────────────────────────────────────────────────────────
output_path = '/Users/ammarali/Downloads/Workspace/Campaign365/Campaign365_User_Guide.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
